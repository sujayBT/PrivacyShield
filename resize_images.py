from docx import Document
import sys
import os
import shutil

mr2_path = os.path.abspath(r"C:\Users\sujay\OneDrive\Documents\MR2_Screenshots_Version.docx")
temp_path = os.path.abspath(r"C:\PrivacyExposureProject\MR2_resized.docx")

shutil.copy2(mr2_path, temp_path)

doc = Document(temp_path)

# Scale all images down by 25% to allow them to fit on the pages better without massive gaps
for shape in doc.inline_shapes:
    shape.width = int(shape.width * 0.75)
    shape.height = int(shape.height * 0.75)

# Also check for empty paragraphs after the target text just in case there are a few
target_idx = -1
for i, p in enumerate(doc.paragraphs):
    if "Privacy Scoring Engine capable" in p.text:
        target_idx = i
        break

if target_idx != -1:
    for i in range(target_idx + 1, min(target_idx + 5, len(doc.paragraphs))):
        p = doc.paragraphs[i]
        has_image = any('<w:drawing>' in r._element.xml or '<w:pict>' in r._element.xml for r in p.runs)
        if not p.text.strip() and not has_image:
            p._element.getparent().remove(p._element)

doc.save(temp_path)

try:
    shutil.copy2(temp_path, mr2_path)
    print("Successfully resized images and removed gaps in MR2_Screenshots_Version.docx")
except Exception as e:
    print(f"Failed to overwrite original (it might be locked by LibreOffice): {e}")
